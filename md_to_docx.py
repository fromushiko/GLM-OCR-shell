from docx import Document

def md_to_docx(input_file, output_file=None):
    with open(input_file, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()
    
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            doc.add_paragraph('')
            continue
        
        if line.startswith('--- 第'):
            doc.add_heading(line.replace('---', '').strip(), level=1)
            continue
        
        if line.startswith('## '):
            doc.add_heading(line[3:], level=2)
            continue
        
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
            continue
        
        doc.add_paragraph(line)
    
    if output_file is None:
        input_path = input_file.replace('.md', '.docx')
        if input_path == input_file:
            input_path = input_file + '.docx'
        output_file = input_path
    
    doc.save(output_file)
    print(f"DOCX生成完成: {output_file}")
    return output_file

if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1:
        input_file = sys.argv[1]
        output_file = sys.argv[2] if len(sys.argv) > 2 else None
        md_to_docx(input_file, output_file)
    else:
        print("用法: python md_to_docx.py <输入md文件> [输出docx文件]")
